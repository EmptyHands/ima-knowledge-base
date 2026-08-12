"""文件解析器测试 - TXT/DOCX 解析 + 短行打包 + 不支持格式拒绝"""
import sys
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

import backend.utils.file_parser as parser_module
from backend.utils.file_parser import parse_file


@pytest.fixture()
def small_pack(monkeypatch):
    """缩小打包目标, 便于在短文本上验证分页"""
    class FakeConfig:
        chunk_size = 60
        chunk_overlap = 4

    monkeypatch.setattr(parser_module, "get_config", lambda: FakeConfig())


def test_parse_txt_short_lines_merge_into_content_pages(small_pack, tmp_path):
    """DEV-002/003 回归: 标题/短行不得单独成碎片页, 应并入相邻内容"""
    f = tmp_path / "doc.txt"
    f.write_text(
        "标题A\n标题B\n" + "内容行" * 30 + "\n"
        "标题C\n" + "另一段内容" * 30,
        encoding="utf-8",
    )
    result = parse_file(str(f))
    assert result["success"] is True
    pages = [p["text"] for p in result["pages"]]
    assert len(pages) == 2
    assert not any(len(t) < 40 for t in pages), f"存在碎片页: {pages}"
    assert any("标题A" in t and "内容行" in t for t in pages)
    assert any("标题C" in t and "另一段内容" in t for t in pages)
    assert [p["page_no"] for p in result["pages"]] == [1, 2]


def test_parse_txt_single_page_when_content_short(tmp_path):
    f = tmp_path / "doc.txt"
    f.write_text("短标题\n一行内容", encoding="utf-8")
    result = parse_file(str(f))
    assert result["success"] is True
    assert len(result["pages"]) == 1
    assert "短标题" in result["pages"][0]["text"]


def test_parse_docx_short_paragraphs_merge(tmp_path):
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("标题段落")
    doc.add_paragraph("正文段落, 内容是完整的说明文字。")
    f = tmp_path / "doc.docx"
    doc.save(str(f))
    result = parse_file(str(f))
    assert result["success"] is True
    joined = "\n".join(p["text"] for p in result["pages"])
    assert "标题段落" in joined and "正文段落" in joined


def test_unsupported_extension(tmp_path):
    f = tmp_path / "doc.exe"
    f.write_bytes(b"MZ")
    result = parse_file(str(f))
    assert result["success"] is False
